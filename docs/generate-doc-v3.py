#!/usr/bin/env python3
"""Генератор: Техническая документация Умная Усадьба v3.0"""
import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═══════════════════════════════════════════
# Хелперы
# ═══════════════════════════════════════════
def H1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    return h

def H2(text):
    h = doc.add_heading(text, level=2)
    return h

def H3(text):
    return doc.add_heading(text, level=3)

def B(text):
    p = doc.add_paragraph(text)
    return p

def BL(text):
    return doc.add_paragraph(text, style='List Bullet')

def BN(text):
    """Bullet Numbered"""
    return doc.add_paragraph(text, style='List Number')

def MONO(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    return p

def CODE(lines):
    """Многострочный моноширинный блок"""
    for line in lines.split('\n'):
        MONO(line)

def TABLE(headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    return t

def LINK(paragraph, url, display_text=None):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    nr = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    nr.append(rPr)
    nr.text = display_text or url
    hl.append(nr)
    paragraph._element.append(hl)

# ═══════════════════════════════════════════
# Титульная страница
# ═══════════════════════════════════════════
T = doc.add_paragraph()
T.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = T.add_run('УМНАЯ УСАДЬБА')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x00, 0xB4, 0xFF)

S = doc.add_paragraph()
S.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = S.add_run('Техническая документация v3.0')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

D = doc.add_paragraph()
D.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = D.add_run('Сервер API · React PWA · Zigbee/MQTT · DuckDB · Docker')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x48, 0x4F, 0x58)

doc.add_paragraph()

M = doc.add_paragraph()
M.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = M.add_run('Июнь 2026 · demossml/smart-estate')
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = RGBColor(0x48, 0x4F, 0x58)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. ОБЩАЯ ИНФОРМАЦИЯ
# ═══════════════════════════════════════════
H1('1. Общая информация')
B('Умная Усадьба — локальная система умного дома на базе Zigbee-устройств. '
  'Сервер написан на TypeScript (Node.js), хранит данные в DuckDB, '
  'общается с устройствами через MQTT (Zigbee2MQTT). Фронтенд — React PWA '
  '(Progressive Web App), устанавливается на телефон как нативное приложение.')

H2('1.1. Ключевые технологии')
TABLE(['Технология', 'Назначение', 'Версия'], [
    ['Node.js', 'Серверная среда', '22.22'],
    ['TypeScript', 'Язык (сервер + фронтенд)', '6.0 / 5.x'],
    ['Express', 'HTTP API сервер', '4.x'],
    ['DuckDB', 'Встраиваемая БД (OLAP)', 'latest'],
    ['MQTT', 'Протокол общения с Zigbee', 'Zigbee2MQTT'],
    ['React', 'Фронтенд-фреймворк', '19.2'],
    ['TailwindCSS', 'CSS-фреймворк', '4.3'],
    ['Vite', 'Сборщик фронтенда', '8.0'],
    ['Docker', 'Контейнеризация сервисов', 'latest'],
    ['Systemd', 'Управление процессами', 'user units'],
])

H2('1.2. Репозиторий')
p = doc.add_paragraph()
p.add_run('GitHub: ').bold = True
LINK(p, 'https://github.com/demossml/smart-estate', 'demossml/smart-estate')
B('Ветка: main (основная)')

H2('1.3. Структура проекта')
MONO('smart-estate/')
MONO('├── server/          # TypeScript сервер (Express + DuckDB + MQTT)')
MONO('│   ├── src/         # Исходный код')
MONO('│   │   ├── index.ts      # Точка входа')
MONO('│   │   ├── api.ts        # REST API (все эндпоинты)')
MONO('│   │   ├── db.ts         # DuckDB схема + хелперы')
MONO('│   │   ├── mqtt-ws.ts    # MQTT клиент + WebSocket')
MONO('│   │   ├── engine.ts     # Движок сценариев')
MONO('│   │   ├── scheduler.ts  # Планировщик по расписанию')
MONO('│   │   ├── demo.ts       # Демо-режим (симуляция)')
MONO('│   │   ├── actions.ts    # Действия сценариев')
MONO('│   │   ├── triggers.ts   # Триггеры сценариев')
MONO('│   │   └── crypto.ts     # Крипто-подпись запросов')
MONO('│   ├── tests/       # Тесты (vitest)')
MONO('│   └── dist/        # Скомпилированный JS')
MONO('├── client-app/      # React PWA фронтенд')
MONO('│   ├── src/')
MONO('│   │   ├── App.tsx           # Главный компонент')
MONO('│   │   ├── main.tsx          # Точка входа React')
MONO('│   │   ├── lib/')
MONO('│   │   │   ├── logger.ts         # Фронтенд-логгер')
MONO('│   │   │   ├── DebugPanel.tsx    # Отладочная панель')
MONO('│   │   │   ├── ErrorBoundary.tsx # Защита от крашей')
MONO('│   │   │   └── utils.ts         # Утилиты')
MONO('│   │   ├── components/ui/   # shadcn-компоненты')
MONO('│   │   └── index.css        # Глобальные стили')
MONO('│   ├── public/')
MONO('│   │   ├── sw.js            # Service Worker')
MONO('│   │   └── manifest.json    # PWA манифест')
MONO('│   └── dist/            # Собранный бандл')
MONO('├── config/          # Конфиги (Mosquitto, Zigbee2MQTT, ACL)')
MONO('├── docker/          # Docker-файлы')
MONO('├── design/          # Дизайн-прототип (HTML + docx)')
MONO('├── docs/            # Документация (wireframes, UX-спека)')
MONO('└── data/            # DuckDB БД (не в гите)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 2. СЕРВЕР (Backend)
# ═══════════════════════════════════════════
H1('2. Сервер API')
H2('2.1. Запуск')
MONO('PORT=8788 \\')
MONO('MQTT_URL=mqtt://localhost:1883 \\')
MONO('node dist/index.js')
B('')
B('Сервер автоматически перезапускается через systemd (user unit). '
  'После падения — рестарт через 5 секунд.')

H2('2.2. Полный список API эндпоинтов')
TABLE(['Метод', 'Путь', 'Описание', 'Rate Limit'], [
    ['GET', '/api/status', 'Статус сервера (устройства, ошибки)', 'Общий'],
    ['GET', '/api/dashboard', 'Агрегированные данные для дашборда', 'Общий'],
    ['GET', '/api/devices', 'Список устройств + телеметрия', 'Общий'],
    ['GET', '/api/devices/:id', 'Детали устройства (история)', 'Общий'],
    ['POST', '/api/devices/:id/on', 'Включить устройство', '30/мин'],
    ['POST', '/api/devices/:id/off', 'Выключить устройство', '30/мин'],
    ['GET', '/api/telemetry', 'История телеметрии (фильтры)', 'Общий'],
    ['GET', '/api/rooms', 'Комнаты + средняя температура', 'Общий'],
    ['GET', '/api/energy', 'Энергопотребление', 'Общий'],
    ['GET', '/api/climate', 'Климат-уставки (термостаты)', 'Общий'],
    ['GET', '/api/climate/:id', 'Уставка конкретного устройства', 'Общий'],
    ['PUT', '/api/climate/:id', 'Обновить уставку (темп/режим)', 'Общий'],
    ['GET', '/api/gates', 'Список ворот/замков', 'Общий'],
    ['POST', '/api/gates/:id/open', 'Открыть ворота', '30/мин'],
    ['POST', '/api/gates/:id/close', 'Закрыть ворота', '30/мин'],
    ['GET', '/api/gates/access-log', 'Журнал доступа', 'Общий'],
    ['GET', '/api/scenarios', 'Сценарии автоматизации', 'Общий'],
    ['POST', '/api/scenarios', 'Создать сценарий', 'Общий'],
    ['PUT', '/api/scenarios/:id', 'Обновить сценарий', 'Общий'],
    ['DELETE', '/api/scenarios/:id', 'Удалить сценарий', 'Общий'],
    ['POST', '/api/scenarios/:id/toggle', 'Вкл/выкл сценарий', '30/мин'],
    ['GET', '/api/scenarios/:id/executions', 'История исполнений', 'Общий'],
    ['GET', '/api/events', 'Последние события (ошибки/команды/смены)', 'Общий'],
    ['GET', '/api/groups', 'Группы устройств', 'Общий'],
    ['GET', '/api/groups/:id', 'Детали группы + участники', 'Общий'],
    ['POST', '/api/groups/:id/add-device', 'Добавить в группу', 'Общий'],
    ['POST', '/api/groups/:id/remove-device', 'Удалить из группы', 'Общий'],
    ['POST', '/api/groups/:id/all-on', 'Включить всю группу', '30/мин'],
    ['POST', '/api/groups/:id/all-off', 'Выключить всю группу', '30/мин'],
    ['GET', '/api/audit', 'Полный аудит (команды/ошибки/смены)', 'Общий'],
    ['GET', '/api/mode', 'Текущий режим (demo/live)', 'Общий'],
    ['POST', '/api/mode', 'Сменить режим', 'Общий'],
    ['POST', '/api/demo/seed', 'Загрузить демо-данные', 'Общий'],
    ['POST', '/api/demo/devices/:id/toggle', 'Тогл демо-устройства', 'Общий'],
    ['POST', '/api/client-logs', 'Принять логи с фронтенда', 'Общий'],
    ['GET', '/api/client-logs', 'Читать логи клиентов', 'Общий'],
    ['GET', '/start', 'React SPA (HTML)', 'Без лимита'],
    ['GET', '/manifest.json', 'PWA манифест', 'Без лимита'],
    ['GET', '/sw.js', 'Service Worker', 'Без лимита'],
    ['GET', '/design', 'Дизайн-прототип', 'Без лимита'],
])

H2('2.3. Rate Limiting')
B('Общий лимит: 500 запросов в минуту. Командные эндпоинты (on/off/open/close/toggle): 30 в минуту. '
  'Статические файлы и страницы — без лимита.')

H2('2.4. Безопасность')
BL('CORS: разрешены localhost и все локальные адреса (192.168.x.x, 10.x.x, 172.16-31.x)')
BL('Helmet: заголовки безопасности (CSP, X-Frame, HSTS и др.)')
BL('Аутентификация: опционально через API_KEYS и X-Signature (крипто-подпись)')
BL('Rate limiting: защита от перебора')

doc.add_page_break()

H2('2.5. База данных (DuckDB)')
B('DuckDB — встраиваемая аналитическая БД. Файл: data/smart-estate.duckdb. '
  'Преимущества: zero-config, быстрые агрегации, SQL-совместимость, нет отдельного сервера.')

H3('Таблицы')
TABLE(['Таблица', 'Назначение', 'Ключевые колонки'], [
    ['devices', 'Реестр устройств', 'ieee_addr, friendly_name, type, room_id, status'],
    ['telemetry', 'Все показания датчиков', 'device_ieee, property, value, unit, ts'],
    ['commands', 'Все отправленные команды', 'device_ieee, command, status, source'],
    ['state_changes', 'История смен состояний', 'device_ieee, old_state, new_state, reason'],
    ['errors', 'Все ошибки системы', 'device_ieee, error_type, error_msg, ts'],
    ['rooms', 'Комнаты', 'id, name, icon'],
    ['scenarios', 'Сценарии автоматизации', 'name, triggers_json, actions_json, active'],
    ['scenario_executions', 'История исполнений', 'scenario_id, actions_fired, success'],
    ['device_groups', 'Группы устройств', 'name, type'],
    ['device_group_members', 'Участники групп', 'group_id, device_ieee'],
    ['climate_setpoints', 'Климат-уставки', 'device_ieee, target_temp, mode, hysteresis'],
    ['gate_access_log', 'Журнал доступа', 'device_ieee, action, source, details'],
])

H3('Индексы')
MONO('idx_telemetry_ts         — по времени телеметрии')
MONO('idx_telemetry_device      — по устройству + свойству + времени')
MONO('idx_commands_device       — по устройству + времени отправки')
MONO('idx_errors_ts             — по времени ошибок')
MONO('idx_state_changes_device  — по устройству + времени')
MONO('idx_scenario_exec_ts      — по времени исполнения')
MONO('idx_scenario_exec_sid     — по ID сценария + времени')

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. MQTT + WebSocket
# ═══════════════════════════════════════════
H1('3. MQTT и WebSocket')
H2('3.1. Архитектура')
B('Сервер подключается к Mosquitto (MQTT брокеру) на localhost:1883. '
  'Zigbee2MQTT публикует данные устройств в топики zigbee2mqtt/#. '
  'Сервер слушает эти топики и пишет телеметрию в DuckDB.')

H3('Поток данных при телеметрии')
BN('Устройство отправляет данные по Zigbee')
BN('Zigbee2MQTT получает их и публикует в MQTT: zigbee2mqtt/0x1234')
BN('Сервер получает сообщение → парсит JSON → извлекает свойства')
BN('Пишет в telemetry (температура, влажность, состояние и т.д.)')
BN('Обновляет devices.last_seen')
BN('Проверяет сценарии (engine.evaluateTelemetry)')
BN('Рассылает всем WebSocket-клиентам')

H2('3.2. Переподключение (экспоненциальный backoff)')
B('При разрыве соединения с MQTT сервер НЕ создаёт дублирующиеся подключения. '
  'Используется экспоненциальный backoff:')
MONO('Попытка 1: 5 секунд')
MONO('Попытка 2: 10 секунд')
MONO('Попытка 3: 20 секунд')
MONO('...')
MONO('Максимум: 60 секунд')
B('')
B('После 3 неудачных попыток сервер переходит в «тихий» режим — '
  'логирует только каждую 10-ю попытку, чтобы не забивать консоль.')

H2('3.3. WebSocket')
B('WebSocket-сервер на path=/ws. При подключении клиента отправляет последние '
  '20 записей телеметрии. Все новые MQTT-сообщения ретранслируются всем '
  'подключённым WebSocket-клиентам в реальном времени.')

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. ФРОНТЕНД (React PWA)
# ═══════════════════════════════════════════
H1('4. Фронтенд — React PWA')
H2('4.1. Экраны приложения')
B('Приложение состоит из 6 экранов (вкладок) + экран установки + отладочная панель:')

TABLE(['Вкладка', 'Иконка', 'Содержание'], [
    ['Главная', '🏠', 'Дашборд: температура, свет, энергия, охрана, климат'],
    ['Устройства', '💡', 'Список устройств сгруппирован по комнатам, фильтр вкл/выкл'],
    ['Климат', '🌡️', 'Термостаты: слайдер целевой температуры, режим (обогрев/охлаждение/авто)'],
    ['Ворота', '🚪', 'Управление воротами (открыть/закрыть) + журнал доступа'],
    ['Сценарии', '⏱️', 'Список сценариев автоматизации, включение/выключение'],
    ['События', '🛡️', 'Лента событий: команды, ошибки, смены состояний'],
])

H2('4.2. Особенности реализации')
H3('PWA (Progressive Web App)')
BL('Устанавливается на домашний экран телефона как нативное приложение')
BL('Service Worker с кешированием (networkFirst для HTML, stale-while-revalidate для ассетов)')
BL('Офлайн-режим: показывает закешированные данные при отсутствии сети')
BL('Иконки: 192×192 и 512×512 px')
BL('Splash screen, theme-color, безопасная зона (safe-area-inset)')

H3('Service Worker (sw.js)')
B('Стратегия кеширования:')
MONO('GET /api/*         → pass-through (не кешируется)')
MONO('GET /assets/*      → stale-while-revalidate (отдаёт кеш, обновляет фон)')
MONO('GET /start         → network-first (всегда свежий HTML, иначе кеш)')
MONO('GET /* (static)    → cache-first (кеш, иначе сеть)')
B('')
B('Версия кеша: usadba-v3. При смене версии старый кеш автоматически очищается.')

H3('Система отладки (Debug Panel)')
BL('Плавающая кнопка 🐛 в правом нижнем углу с счётчиком ошибок')
BL('Автоматически логирует: ошибки JS, предупреждения, долгие кадры (>100ms), медленные fetch (>3s)')
BL('Кнопка «Грише» — отправляет логи на сервер (POST /api/client-logs)')
BL('ErrorBoundary — перехватывает краши React и показывает кнопку «Перезагрузить»')
BL('Логи сохраняются в sessionStorage для выживания между перезагрузками')

H3('Оптимизации производительности')
BL('Слайдер климата: дебаунс 400ms — API-запрос только после отпускания, не на каждом пикселе')
BL('Тоглы: оптимистичный UI, рефреш данных с задержкой 300ms для плавной анимации')
BL('Тосты: pointer-events-none — не блокируют тапы по интерфейсу')
BL('Скролл: -webkit-overflow-scrolling: touch + overscroll-behavior: contain')
BL('Полинг: 8 запросов каждые 5 секунд (статус, дашборд, устройства, климат, ворота, логи, сценарии, события)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 5. ДИЗАЙН-ПРОЕКТ
# ═══════════════════════════════════════════
H1('5. Дизайн-проект')
H2('5.1. Дизайн-токены')
TABLE(['Токен', 'Значение', 'Использование'], [
    ['--bg-primary', '#0D1117', 'Фон страницы'],
    ['--bg-secondary', '#161B22', 'Фон карточек, хедера, таб-бара'],
    ['--bg-tertiary', '#21262D', 'Фон элементов ввода, скелетонов'],
    ['--text-primary', '#E6EDF3', 'Основной текст'],
    ['--text-secondary', '#8B949E', 'Второстепенный текст'],
    ['--text-tertiary', '#484F58', 'Метки, временные метки'],
    ['--accent-blue', '#00B4FF', 'Акцентный (кнопки, активные элементы)'],
    ['--accent-green', '#00FF9D', 'Успех, онлайн, подтверждение'],
    ['--accent-orange', '#FF8C00', 'Предупреждения, нагрев'],
    ['--accent-red', '#FF3860', 'Ошибки, тревога, опасность'],
    ['--accent-yellow', '#FFD700', 'Золотой акцент'],
    ['--accent-purple', '#B060FF', 'Датчики движения'],
    ['--border-subtle', 'rgba(255,255,255,0.06)', 'Границы карточек'],
    ['--border-active', 'rgba(0,180,255,0.25)', 'Активные границы'],
    ['--radius-card', '16px', 'Скругление карточек'],
    ['--radius-btn', '12px', 'Скругление кнопок'],
    ['--radius-tile', '12px', 'Скругление плиток устройств'],
])

H2('5.2. Типографика')
B('Шрифт: системный стек (-apple-system, BlinkMacSystemFont, Segoe UI, Inter). '
  'Моноширинный: SF Mono / JetBrains Mono (для данных). '
  'Базовый размер: 15px, межстрочный: 1.4.')

H2('5.3. Адаптивность')
BL('Мобильный (≤480px): основной вид, max-width контейнера 480px')
BL('Планшет (≥768px): #root max-width: 100%, grid: 3 колонки')
BL('Поддержка поворота экрана (resize listener в логгере)')

H2('5.4. Прототип')
p = doc.add_paragraph()
p.add_run('HTML-прототип: ').bold = True
LINK(p, 'http://localhost:8788/design', '/design (доступен при запущенном сервере)')
B('Дизайн-документы в формате .docx:')
BL('design/Дизайн-проект_Умная_Усадьба_v1.0.docx')
BL('docs/Wireframes_Умная_Усадьба.docx')
BL('docs/Дизайн-проект_Умная_Усадьба_v2.0.docx')
BL('docs/Умная Усадьба - UX Спецификация.docx')

doc.add_page_break()

# ═══════════════════════════════════════════
# 6. СЦЕНАРИИ АВТОМАТИЗАЦИИ
# ═══════════════════════════════════════════
H1('6. Сценарии автоматизации')
B('Всего 10 сценариев. Каждый имеет триггеры (JSON), действия (JSON), '
  'опциональное расписание (cron или sunset).')

TABLE(['ID', 'Название', 'Триггер', 'Действие', 'Расписание'], [
    ['1', 'Вентиляция по CO₂', 'CO₂ > 1000 ppm', 'Открыть клапан вентиляции', '—'],
    ['2', 'Свет при закате', 'Освещённость < 50 lux', 'Включить садовый свет', 'Закат −30 мин'],
    ['3', 'Охрана периметра', 'Дверь открыта', 'Уведомление', '—'],
    ['4', 'Полив по расписанию', 'Влажность почвы < 40%', 'Открыть клапан полива на 10 мин', '0 6 * * *'],
    ['5', 'Обогрев при низкой T°', 'Температура < 18°C', 'Включить котёл', '—'],
    ['6', 'Защита от протечки', 'Обнаружена протечка', 'Перекрыть воду', '—'],
    ['7', 'Отключение света', '00:00', 'Весь свет OFF', '0 0 * * *'],
    ['8', 'Эко-режим', 'Все ушли', 'Снизить темп., выкл. свет', '—'],
    ['9', 'Гостевое освещение', 'Движение в коридоре ночью', 'Тёплый свет 30%', '22:00–06:00'],
    ['10', 'Шторм-режим', 'Скорость ветра > 15 м/с', 'Закрыть ворота, маркизы', '—'],
])

H2('6.1. Движок сценариев (engine.ts)')
BL('При каждом новом показании датчика вызывает evaluateTelemetry()')
BL('Проверяет триггеры всех активных сценариев')
BL('При совпадении всех условий (логика ALL/ANY) — выполняет действия')
BL('Пишет результат в scenario_executions')

H2('6.2. Планировщик (scheduler.ts)')
BL('Проверяет сценарии с schedule_json (cron или sunset)')
BL('Запускается каждую минуту')
BL('При совпадении расписания + условий — выполняет действия')

doc.add_page_break()

# ═══════════════════════════════════════════
# 7. РАЗВЁРТЫВАНИЕ
# ═══════════════════════════════════════════
H1('7. Развёртывание и DevOps')
H2('7.1. Systemd (продакшен)')
B('Сервер запускается как user systemd unit:')
MONO('systemctl --user status smart-estate.service')
B('')
B('Конфигурация: ~/.config/systemd/user/smart-estate.service')
MONO('[Service]')
MONO('Type=simple')
MONO('ExecStart=/home/admingimolost/.local/bin/node server/dist/index.js')
MONO('WorkingDirectory=~/smart-estate/server')
MONO('Environment=PORT=8788')
MONO('Environment=MQTT_URL=mqtt://localhost:1883')
MONO('Restart=always')
MONO('RestartSec=5')

H2('7.2. Docker (разработка)')
B('docker-compose.yml запускает Mosquitto + Zigbee2MQTT.')
MONO('cd ~/smart-estate && docker compose up -d')
B('')
MONO('Сервисы:')
MONO('  mosquitto       — MQTT брокер (порт 1883)')
MONO('  zigbee2mqtt     — Zigbee → MQTT мост')

H2('7.3. Сборка и деплой')
H3('Сервер')
BN('cd server && npx tsc — компиляция TypeScript')
BN('systemctl --user restart smart-estate.service — рестарт')

H3('Фронтенд')
BN('cd client-app && npm run build — сборка Vite (tsc + vite build)')
BN('Статика раздаётся сервером из client-app/dist/')

H3('Полный цикл обновления')
BN('cd client-app && npm run build')
BN('systemctl --user restart smart-estate.service')
BN('На телефоне: жёсткая перезагрузка PWA (смахнуть и открыть заново, или 2 рефреша подряд)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 8. ИСТОРИЯ ИЗМЕНЕНИЙ
# ═══════════════════════════════════════════
H1('8. История изменений (Changelog)')

H2('v3.0 — Июнь 2026 (текущая)')
B('Крупное обновление: React PWA + отладка + исправления.')
BL('React PWA фронтенд (client-app): TailwindCSS, shadcn-стиль, 6 экранов')
BL('Дебаг-панель: логгер ошибок, jank-детектор, кнопка «Грише»')
BL('ErrorBoundary: защита от крашей React')
BL('MQTT fix: устранена утечка памяти (2GB → 64MB), экспоненциальный backoff')
BL('Rate limit: 120 → 500 req/min')
BL('CORS: разрешены все локальные адреса')
BL('Cache-Control: no-store для HTML/SW/manifest')
BL('POST/GET /api/client-logs — удалённый приём логов')
BL('Слайдер климата: дебаунс 400ms, локальное состояние')
BL('Тоглы: задержка рефреша 300ms для плавной анимации')
BL('Тосты: pointer-events-none')
BL('TypeScript: paths alias (@/) исправлен')

H2('v2.0 — Май 2026')
BL('/api/dashboard — агрегированный эндпоинт для фронтенда')
BL('Demo-режим: симуляция датчиков для тестирования')
BL('POST /api/mode — переключение demo ↔ live')
BL('Сервер слушает на 0.0.0.0 (доступ с телефона)')
BL('Автоопределение IP для доступа')

H2('v1.0 — Май 2026 (инициализация)')
BL('TypeScript сервер: Express + DuckDB')
BL('MQTT интеграция: подписка zigbee2mqtt/#')
BL('REST API: устройства, телеметрия, комнаты, энергия')
BL('Docker: Mosquitto + Zigbee2MQTT')
BL('Тесты: vitest (12 файлов)')
BL('Дизайн-документы: UX-спецификация, wireframes, прототип')

doc.add_page_break()

# ═══════════════════════════════════════════
# 9. ТЕКУЩЕЕ СОСТОЯНИЕ
# ═══════════════════════════════════════════
H1('9. Текущее состояние (Июнь 2026)')

H2('9.1. Что работает')
BL('✅ Сервер API — все 37 эндпоинтов')
BL('✅ DuckDB — 16 устройств, 6 климат-уставок, 2 ворот, 10 сценариев')
BL('✅ React PWA — 6 экранов, установка на телефон')
BL('✅ Service Worker — кеширование, networkFirst для HTML')
BL('✅ Дебаг-панель — логгер, отправка логов на сервер')
BL('✅ MQTT — стабильное переподключение, нет утечек')
BL('✅ Rate limiting — 500 req/min')
BL('✅ CORS — все локальные адреса')
BL('✅ Demo-режим — симуляция датчиков')
BL('✅ WebSocket — ретрансляция в реальном времени')

H2('9.2. В процессе / запланировано')
BL('⬜ Интеграция с Telegram Mini App')
BL('⬜ Push-уведомления о тревогах')
BL('⬜ Графики телеметрии (история температуры, энергии)')
BL('⬜ Голосовое управление')
BL('⬜ Автоматическое обновление PWA (баннер «доступна новая версия»)')

H2('9.3. Известные ограничения')
BL('Zigbee2MQTT требует USB-стик (Coordinator), не эмулируется')
BL('MQTT брокер не обязателен для работы сервера, но без него нет real-time данных')
BL('Service Worker при первом деплое может требовать двойной перезагрузки (известная особенность SW lifecycle)')

doc.add_page_break()

# ═══════════════════════════════════════════
# 10. ЧЕКЛИСТ РАЗРАБОТЧИКА
# ═══════════════════════════════════════════
H1('10. Чеклист разработчика')

H2('10.1. Быстрый старт')
BN('Клонировать репо: git clone https://github.com/demossml/smart-estate')
BN('Сервер: cd server && npm install && npx tsc')
BN('Фронтенд: cd client-app && npm install && npm run build')
BN('Запустить: systemctl --user start smart-estate.service')
BN('Проверить: curl http://localhost:8788/api/status')

H2('10.2. После изменений в коде')
BN('Сервер: npx tsc && systemctl --user restart smart-estate.service')
BN('Фронтенд: npm run build && systemctl --user restart smart-estate.service')
BN('Проверить: curl http://localhost:8788/api/status → {"ok":true}')
BN('Проверить ассеты: curl -sI http://localhost:8788/assets/index-*.js → 200')
BN('На телефоне: жёсткая перезагрузка PWA')

H2('10.3. Типичные проблемы и решения')
TABLE(['Проблема', 'Причина', 'Решение'], [
    ['Белый экран после деплоя', 'Старый SW кеширует HTML с мёртвой ссылкой на JS', 'Перезагрузить 2 раза или удалить PWA и добавить заново'],
    ['Сервер не отвечает', 'MQTT flood → утечка памяти', 'Проверить mqtt-ws.ts — не должно быть множественных connect()'],
    ['Слайдер дёргается', 'API-запрос на каждом пикселе', 'Должен быть дебаунс 400ms (исправлено в v3.0)'],
    ['CORS ошибка на телефоне', 'Origin не в белом списке', 'Проверить CORS в api.ts — должны быть все локальные адреса'],
    ['Rate limit ошибка', '8 polls/min × 5s + действия > лимита', 'Лимит поднят до 500 (исправлено в v3.0)'],
])

H2('10.4. Полезные команды')
MONO('# Логи сервера')
MONO('journalctl --user -u smart-estate.service -f')
MONO('')
MONO('# Статус')
MONO('systemctl --user status smart-estate.service')
MONO('')
MONO('# Рестарт')
MONO('systemctl --user restart smart-estate.service')
MONO('')
MONO('# Проверка всех API')
MONO('curl -s http://localhost:8788/api/status | jq .')
MONO('curl -s http://localhost:8788/api/devices | jq ".devices | length"')
MONO('')
MONO('# Клиентские логи (удалённо)')
MONO('curl -s http://localhost:8788/api/client-logs | jq .')
MONO('')
MONO('# Сборка фронтенда')
MONO('cd client-app && npm run build')
MONO('')
MONO('# Git')
MONO('git add -A && git commit -m "..." && git push origin main')

# ═══════════════════════════════════════════
# Сохранение
# ═══════════════════════════════════════════
output = '/home/admingimolost/smart-estate/docs/Техническая_документация_Умная_Усадьба_v3.0.docx'
doc.save(output)
print(f'✅ Saved: {output}')
print(f'Pages: ~15-20 (A4), Sections: 10')
